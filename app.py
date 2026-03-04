import streamlit as st
import mne
import plotly.graph_objects as go
import os
import numpy as np
import pandas as pd
from openai import OpenAI
from datetime import datetime
from docx import Document
from io import BytesIO

# --- 1. 安全配置 (优先读取云端 Secrets) ---
try:
    DEEPSEEK_API_KEY = st.secrets["DEEPSEEK_API_KEY"]
except:
    DEEPSEEK_API_KEY = "sk-xxxxxxxxxxxxxxxx" # 仅供本地临时测试

BASE_URL = "https://api.deepseek.com"

# --- 页面 UI 风格 ---
st.set_page_config(page_title="思言的脑电工作台", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
    <style>
    .main { background-color: #fcfcfc; }
    .stButton>button { border-radius: 8px; font-weight: 500; height: 3em; width: 100%; }
    div[data-testid="stExpander"] { border: 1px solid #f0f2f6; box-shadow: 0 4px 12px rgba(0,0,0,0.03); background: white; border-radius: 12px; }
    .report-box { padding: 20px; background-color: #f9f9f9; border-radius: 10px; border-left: 5px solid #4a90e2; }
    </style>
    """, unsafe_allow_html=True)

# 状态初始化
if "messages" not in st.session_state: st.session_state.messages = []
if "raw_obj" not in st.session_state: st.session_state.raw_obj = None
if "cleaned_raw" not in st.session_state: st.session_state.cleaned_raw = None
if "report_content" not in st.session_state: st.session_state.report_content = ""

# --- 2. 侧边栏：资源管理 ---
with st.sidebar:
    st.header("📂 资源管理")
    uploaded_files = st.file_uploader("导入脑电数据", type=["vhdr", "vmrk", "eeg", "edf", "set", "fif"], accept_multiple_files=True)
    st.divider()
    st.warning("💡 大文件上传请务必开启下方开关")
    do_resample = st.toggle("极速处理模式 (自动降采样)", value=True)
    target_fs = st.number_input("目标采样率 (Hz)", value=250)

TEMP_DIR = "temp_eeg_processing"
if not os.path.exists(TEMP_DIR): os.makedirs(TEMP_DIR)

# --- 3. 核心导出函数 ---
def create_word_report(content):
    doc = Document()
    doc.add_heading('思言脑电工作台 - 实验报告', 0)
    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_to_excel(raw):
    # 仅导出前 15 秒数据作为样本，防止内存溢出
    data, times = raw[:, :int(raw.info['sfreq'] * 15)]
    df = pd.DataFrame(data.T, columns=raw.ch_names)
    df.insert(0, 'Time(s)', times)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    return output.getvalue()

# --- 4. 主界面逻辑 ---
st.title("✨ 思言的脑电工作台")

if uploaded_files:
    main_file = None
    for f in uploaded_files:
        path = os.path.join(TEMP_DIR, f.name)
        if not os.path.exists(path):
            with open(path, "wb") as b: b.write(f.getbuffer())
        if f.name.lower().endswith(('.vhdr', '.edf', '.set', '.fif')): 
            main_file = path

    # 懒加载：初次读取不 preload
    if main_file and st.session_state.raw_obj is None:
        with st.spinner("正在解析文件结构..."):
            st.session_state.raw_obj = mne.io.read_raw(main_file, preload=False)

if st.session_state.raw_obj:
    raw = st.session_state.raw_obj
    
    with st.expander("🛠️ 数据处理与多格式导出", expanded=True):
        c1, c2, c3 = st.columns(3)
        
        with c1:
            st.write("**核心操作**")
            if st.button("🚀 执行全自动清洗"):
                with st.status("正在进行内存优化处理...") as status:
                    # 此时才真正加载数据，并立即降采样
                    status.write("正在加载数据并执行降采样...")
                    curr_raw = raw.copy().load_data()
                    if do_resample and curr_raw.info['sfreq'] > target_fs:
                        curr_raw.resample(target_fs)
                    
                    status.write("正在执行滤波与坏道插值...")
                    curr_raw.filter(0.1, 45.0, verbose=False)
                    
                    # 简单坏道识别
                    data_tmp = curr_raw.get_data()
                    stds = np.std(data_tmp, axis=1)
                    bads = [curr_raw.ch_names[i] for i, s in enumerate(stds) if s > 4 * np.median(stds) or s < 1e-7]
                    if bads:
                        curr_raw.info['bads'] = bads
                        curr_raw.interpolate_bads(reset_bads=True)
                    
                    status.write("正在运行 ICA 伪迹剔除...")
                    ica = mne.preprocessing.ICA(n_components=min(len(curr_raw.ch_names)-1, 10), random_state=97)
                    ica.fit(curr_raw, verbose=False)
                    ica.exclude = [0, 1]
                    st.session_state.cleaned_raw = ica.apply(curr_raw, verbose=False)
                    status.update(label="✅ 清洗完成！", state="complete")

            if st.button("📝 生成 AI 实验报告"):
                if st.session_state.cleaned_raw:
                    with st.spinner("思言正在思考并撰写..."):
                        client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=BASE_URL)
                        prompt = f"当前数据有{len(st.session_state.cleaned_raw.ch_names)}个通道。请撰写一段专业的脑电预处理实验描述。"
                        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": prompt}])
                        st.session_state.report_content = res.choices[0].message.content
                else:
                    st.warning("请先执行清洗")

        with c2:
            st.write("**数据导出**")
            if st.session_state.cleaned_raw:
                st.session_state.cleaned_raw.save("temp_clean.fif", overwrite=True, verbose=False)
                with open("temp_clean.fif", "rb") as f:
                    st.download_button("💾 下载清洗数据 (FIF)", f, "clean_data.fif")
                
                excel_file = export_to_excel(st.session_state.cleaned_raw)
                st.download_button("📊 导出 Excel (15s采样)", excel_file, "eeg_sample.xlsx")

        with c3:
            st.write("**报告导出**")
            if st.session_state.report_content:
                st.download_button("📄 导出 Markdown", st.session_state.report_content, "report.md")
                word_file = create_word_report(st.session_state.report_content)
                st.download_button("📘 导出 Word (.docx)", word_file, "report.docx")

    # --- 5. 交互式可视化看板 ---
    st.divider()
    st.subheader("📊 数据可视化看板")
    display_raw = st.session_state.cleaned_raw if st.session_state.cleaned_raw else st.session_state.raw_obj
    
    # 限制时域预览长度，保护浏览器内存
    start_sec = st.slider("时间轴预览定位 (秒)", 0.0, float(display_raw.times[-1]-5), 0.0)
    idx_s, idx_e = display_raw.time_as_index([start_sec, start_sec + 5])
    
    # 仅提取这 5 秒的数据进行绘图
    plot_data = display_raw.get_data(start=idx_s, stop=idx_e)
    plot_times = display_raw.times[idx_s:idx_e]
    
    wave_fig = go.Figure()
    for i in range(min(12, len(display_raw.ch_names))):
        wave_fig.add_trace(go.Scatter(x=plot_times, y=plot_data[i] + (i * 200e-6), name=display_raw.ch_names[i], line=dict(width=1)))
    wave_fig.update_layout(height=500, margin=dict(l=10, r=10, t=10, b=10), template="plotly_white", dragmode="pan", xaxis_title="Time (s)")
    st.plotly_chart(wave_fig, use_container_width=True, config={'scrollZoom': True})

    # PSD 分析
    psds, freqs = display_raw.compute_psd(fmax=50, verbose=False).get_data(return_freqs=True)
    psd_fig = go.Figure()
    for i in range(min(8, len(display_raw.ch_names))):
        psd_fig.add_trace(go.Scatter(x=freqs, y=10*np.log10(psds[i]), name=display_raw.ch_names[i]))
    psd_fig.update_layout(height=400, margin=dict(l=10, r=10, t=10, b=10), template="plotly_white", xaxis_title="Frequency (Hz)", yaxis_title="dB")
    st.plotly_chart(psd_fig, use_container_width=True)

    # --- 6. 问我 ---
    st.divider()
    st.subheader("💬 问我")
    chat_box = st.container(height=350)
    for msg in st.session_state.messages:
        with chat_box:
            with st.chat_message(msg["role"]): st.markdown(msg["content"])
    
    if prompt := st.chat_input("询问关于实验分析的问题..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with chat_box:
            with st.chat_message("user"): st.markdown(prompt)
            with st.chat_message("assistant"):
                try:
                    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=BASE_URL)
                    res = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[{"role": "system", "content": "你是脑电专家助手思言。"}, 
                                  {"role": "user", "content": prompt}]
                    )
                    st.markdown(res.choices[0].message.content)
                    st.session_state.messages.append({"role": "assistant", "content": res.choices[0].message.content})
                except Exception as e: st.error(f"对话出现波动: {e}")
else:
    st.info("👋 欢迎来到思言的工作台。请在侧边栏上传数据开始。")